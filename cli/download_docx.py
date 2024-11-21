import io

import boto3
from docx import Document


def import_docx_from_s3(bucket_name, object_key):
    """
    Download a .docx file from S3 and read its content

    Returns:
        str: Content of the .docx file
    """

    try:
        s3_client = boto3.client("s3")
        response = s3_client.get_object(Bucket=bucket_name, Key=object_key)
        docx_file = io.BytesIO(response["Body"].read())
        doc = Document(docx_file)
        return doc
    except Exception as e:
        raise Exception(f"Error processing the document: {e}")


def import_docx_from_local_file(filepath):
    """
    Read a .docx file from a local file and return its content as a string

    Returns:
        str: Content of the .docx file
    """
    try:
        doc = Document(filepath)
        return doc

    except Exception as e:
        raise Exception(f"Error processing the document: {e}")


def prepare_string_from_docx(doc: Document):
    """
    Read a .docx file directly from S3 and return its content as a string

    Args:
        bucket_name (str): Name of the S3 bucket
        object_key (str): Path to the object in the bucket

    Returns:
        str: Content of the .docx file
    """
    try:

        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                full_text.append(para.text)

        # Join all paragraphs with newlines
        return "\n".join(full_text)

    except boto3.exceptions.BotoResponseError as e:
        raise Exception(f"Failed to read from S3: {e}")
    except Exception as e:
        raise Exception(f"Error processing the document: {e}")
